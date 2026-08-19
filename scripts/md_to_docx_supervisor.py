import os
import re
import sys

sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC_DIR = r"D:\aprecap\docs\markdown_cursos\3_Supervisor_de_Seguridad"
OUT_PATH = r"C:\Users\Vickoto\Desktop\Curso_Supervisor_Consolidado_v2.docx"

SUBMODULOS = [
    "Modulo_1.1_Contrato_de_Trabajo_y_Jornada_Laboral.md",
    "Modulo_1.2_Reglamento_209_Ley_21659.md",
    "Modulo_1.3_Ley_21659_Seguridad_Privada_en_la_Practica.md",
    "Modulo_1.4_Derecho_Penal_y_Detencion.md",
    "Modulo_1.5_Derechos_Humanos_Uso_de_la_Fuerza_y_Datos_Personales.md",
    "Modulo_1.6_Ley_16744_Accidentes_y_Enfermedades.md",
    "Modulo_1.7_Decreto_594_Higiene_y_Seguridad.md",
    "Modulo_2.1_Prevencion_de_Riesgos_en_el_Puesto.md",
    "Modulo_2.2_Control_de_Incendios_y_Emergencias.md",
    "Modulo_3.1_Directivas_de_Funcionamiento_y_OS10.md",
    "Modulo_3.2_Estudios_de_Seguridad_y_Pautas_de_Puesto.md",
    "Modulo_4.1_Liderazgo_y_Supervision_de_Equipos.md",
    "Modulo_4.2_Resolucion_de_Conflictos_y_Ley_Karin.md",
    "Modulo_5.1_Sistemas_de_Alarma_y_Monitoreo.md",
    "Modulo_5.2_Comunicacion_y_Enlace.md",
    "Modulo_6.1_Eventos_Masivos_Ley_21659.md",
    "Modulo_6.2_Registros_Operativos_e_Informes.md",
    "Modulo_6.3_Manejo_de_Incidentes_del_Supervisor.md",
]


def add_runs(par, text):
    """Agrega runs a un párrafo interpretando **negritas**."""
    for i, seg in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if seg == "":
            continue
        run = par.add_run(seg)
        if i % 2 == 1:
            run.bold = True


def leer_titulo(texto):
    """Devuelve el título sin el prefijo '# Submódulo X.Y:' cuando existe."""
    return re.sub(r"^#\s*", "", texto).strip()


def main():
    doc = Document()

    # Portada
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("APRECAP · Capacitaciones y Asesorías")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x00, 0x21, 0x59)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Curso de Supervisor de Seguridad Privada")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(0x00, 0x21, 0x59)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Módulos de Estudio — Documento Consolidado para Revisión")
    r.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("18 de agosto de 2026")
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Este documento reúne el contenido completo de los 18 submódulos del curso, "
        "con la normativa chilena vigente, para su revisión y comentarios."
    )
    r.italic = True
    r.font.size = Pt(11)

    doc.add_page_break()

    # Índice
    doc.add_heading("Índice de submódulos", level=1)
    for archivo in SUBMODULOS:
        ruta = os.path.join(SRC_DIR, archivo)
        texto = open(ruta, encoding="utf-8").read()
        primera = [l for l in texto.split("\n") if l.startswith("# ")][0]
        titulo = leer_titulo(primera)
        par = doc.add_paragraph()
        add_runs(par, titulo)
        par.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # Contenido
    for archivo in SUBMODULOS:
        ruta = os.path.join(SRC_DIR, archivo)
        lineas = open(ruta, encoding="utf-8").read().split("\n")
        for linea in lineas:
            s = linea.rstrip()
            if not s.strip():
                continue
            if s.startswith("# "):
                doc.add_heading(leer_titulo(s), level=1)
            elif s.startswith("## "):
                doc.add_heading(s[3:].strip(), level=2)
            elif s.startswith("### "):
                doc.add_heading(s[4:].strip(), level=3)
            elif s.startswith("- "):
                par = doc.add_paragraph(style="List Bullet")
                add_runs(par, s[2:].strip())
            elif re.match(r"^\d+\.\s", s):
                par = doc.add_paragraph()
                add_runs(par, s.strip())
            else:
                par = doc.add_paragraph()
                add_runs(par, s.strip())

    doc.save(OUT_PATH)
    print("[OK]", OUT_PATH)


if __name__ == "__main__":
    main()
